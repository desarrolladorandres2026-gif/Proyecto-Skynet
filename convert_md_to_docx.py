#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Convertir archivos Markdown a Word (.docx) manteniendo estructura y formato.
Recorre recursivamente la carpeta 'docs' y genera .docx en la misma ubicación.
"""

import os
import sys
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from markdown2 import markdown

# Fix para Windows encoding
if sys.platform == 'win32':
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

def extract_title_from_markdown(content):
    """Extrae el título del primer H1 en el markdown."""
    match = re.search(r'^#\s+(.+)$', content, re.MULTILINE)
    return match.group(1).strip() if match else "Documento"

def parse_markdown_to_docx(markdown_content, docx_path):
    """
    Convierte markdown a .docx con estilos básicos pero presentables.
    """
    doc = Document()

    # Título principal
    title = extract_title_from_markdown(markdown_content)
    title_para = doc.add_paragraph(title)
    title_para.style = 'Heading 1'
    title_run = title_para.runs[0]
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)  # Azul oscuro

    doc.add_paragraph()  # Espaciador

    # Procesar línea por línea
    lines = markdown_content.split('\n')
    current_list_type = None
    list_level = 0

    for line in lines:
        line = line.rstrip()

        # Headers
        if line.startswith('# '):
            p = doc.add_paragraph(line[2:], style='Heading 1')
            p.runs[0].font.size = Pt(18)
        elif line.startswith('## '):
            p = doc.add_paragraph(line[3:], style='Heading 2')
            p.runs[0].font.size = Pt(14)
        elif line.startswith('### '):
            p = doc.add_paragraph(line[4:], style='Heading 3')
            p.runs[0].font.size = Pt(12)
        elif line.startswith('#### '):
            p = doc.add_paragraph(line[5:], style='Heading 4')

        # Listas no-ordenadas
        elif line.startswith('- '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
            list_level = 0
        elif line.startswith('  - '):
            p = doc.add_paragraph(line[4:], style='List Bullet 2')

        # Listas ordenadas
        elif re.match(r'^\d+\.\s', line):
            match = re.match(r'^(\d+)\.\s(.+)$', line)
            if match:
                p = doc.add_paragraph(match.group(2), style='List Number')

        # Código monobloque
        elif line.startswith('```'):
            continue  # Se ignoran las líneas de cierre de bloques

        # Líneas vacías
        elif not line.strip():
            continue

        # Texto normal y con formato
        else:
            if line.strip():
                p = doc.add_paragraph(line)
                # Procesar **bold** e *italic*
                for run in p.runs:
                    if '**' in run.text:
                        run.bold = True
                    if '*' in run.text and '**' not in run.text:
                        run.italic = True

    doc.save(docx_path)

def convert_docs_folder():
    """Busca todos los .md en 'docs' y los convierte a .docx"""
    docs_folder = Path('docs')

    if not docs_folder.exists():
        print("[ERROR] Carpeta 'docs' no encontrada")
        return

    md_files = list(docs_folder.rglob('*.md'))

    if not md_files:
        print("[WARN] No se encontraron archivos .md en la carpeta docs/")
        return

    print(f"[INFO] Encontrados {len(md_files)} archivos .md\n")

    for md_file in sorted(md_files):
        try:
            # Leer el markdown
            with open(md_file, 'r', encoding='utf-8') as f:
                content = f.read()

            # Ruta de salida: mismo nombre pero con extensión .docx
            docx_file = md_file.with_suffix('.docx')

            # Convertir
            parse_markdown_to_docx(content, str(docx_file))
            print(f"[OK] {md_file} -> {docx_file}")

        except Exception as e:
            print(f"[ERROR] {md_file}: {str(e)}")

    print(f"\n[DONE] Conversion completada. Se generaron {len(md_files)} archivos .docx")

if __name__ == '__main__':
    convert_docs_folder()
